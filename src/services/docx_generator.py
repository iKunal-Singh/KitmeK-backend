"""
DOCX Generator — Converts validated lesson JSON into professional DOCX.

Uses python-docx to produce styled output per Architecture Doc Section 4.3.4.
Handles ALL missing optional fields gracefully.
"""

import io
import logging
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

TITLE_COLOR = RGBColor(0x1F, 0x4E, 0x79)  # #1F4E79
SECTION_COLOR = RGBColor(0x2E, 0x75, 0xB6)  # Section heading accent
LABEL_COLOR = RGBColor(0x59, 0x56, 0x59)  # Gray for labels


def _safe_str(value: Any, default: str = "") -> str:
    """Safely convert value to string."""
    if value is None:
        return default
    if isinstance(value, str):
        return value
    return str(value)


def _add_styled_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading with styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = TITLE_COLOR
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = SECTION_COLOR
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = SECTION_COLOR
            run.font.size = Pt(13)


def _add_label_value(doc: Document, label: str, value: str) -> None:
    """Add a bold label followed by a value."""
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.color.rgb = LABEL_COLOR
    run_label.font.size = Pt(10)
    run_value = para.add_run(value)
    run_value.font.size = Pt(10)


def _add_narration_line(doc: Document, label: str, text: str) -> None:
    """Add a narration line with label."""
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = para.add_run(text)
    run_text.font.size = Pt(10)
    run_text.italic = True


class DocxGenerator:
    """Generates styled DOCX from validated lesson data."""

    def generate(
        self,
        lesson_data: dict[str, Any],
        grade: str = "",
        subject: str = "",
        topic_name: str = "",
        chapter_name: str = "",
        validation_report: Optional[dict[str, Any]] = None,
    ) -> bytes:
        """
        Generate a DOCX document from lesson data.

        Returns bytes of the DOCX file.
        Handles all missing optional fields gracefully.
        """
        try:
            doc = Document()
            self._set_document_defaults(doc)
            self._add_title_section(doc, topic_name, grade, subject, chapter_name)
            self._add_learning_objective(doc, lesson_data)
            self._add_opening_narration(doc, lesson_data)
            self._add_on_screen_opening(doc, lesson_data)
            self._add_narrated_explanation(doc, lesson_data)
            self._add_interactive_activity(doc, lesson_data)
            self._add_doubts_discussion(doc, lesson_data)
            self._add_quick_quiz(doc, lesson_data)
            self._add_conclusion(doc, lesson_data)
            if validation_report:
                self._add_validation_appendix(doc, validation_report)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()

        except Exception as exc:
            logger.error("DOCX generation failed: %s", exc)
            raise RuntimeError(f"DOCX generation failed: {exc}") from exc

    def _set_document_defaults(self, doc: Document) -> None:
        """Set default styles for the document."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_title_section(
        self, doc: Document, topic_name: str, grade: str, subject: str, chapter_name: str
    ) -> None:
        """Add title page / header section."""
        # Title — Heading 1, centered, blue
        title_para = doc.add_heading(topic_name or "Untitled Lesson", level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = TITLE_COLOR
            run.font.size = Pt(24)

        # Subtitle — grade/subject/chapter
        subtitle_parts = []
        if grade:
            subtitle_parts.append(f"Grade {grade}")
        if subject:
            subtitle_parts.append(f"Subject: {subject}")
        if chapter_name:
            subtitle_parts.append(f"Chapter: {chapter_name}")

        if subtitle_parts:
            subtitle = doc.add_heading(" | ".join(subtitle_parts), level=2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in subtitle.runs:
                run.font.color.rgb = SECTION_COLOR
                run.font.size = Pt(14)

        doc.add_paragraph("")  # spacing

    def _add_learning_objective(self, doc: Document, lesson_data: dict) -> None:
        """Add Learning Objective section."""
        objective = _safe_str(lesson_data.get("learning_objective"))
        if not objective:
            return
        _add_styled_heading(doc, "Learning Objective", level=2)
        para = doc.add_paragraph(objective)
        para.style.font.size = Pt(11)

    def _add_opening_narration(self, doc: Document, lesson_data: dict) -> None:
        """Add Opening Narration section with 4 lines."""
        opening = lesson_data.get("opening_narration")
        if not opening:
            return

        _add_styled_heading(doc, "Opening Narration", level=2)

        if isinstance(opening, dict):
            for key in sorted(opening.keys()):
                line = _safe_str(opening[key])
                if line:
                    _add_narration_line(doc, f"Teacher says ({key})", line)
        elif isinstance(opening, str):
            _add_narration_line(doc, "Teacher says", opening)

    def _add_on_screen_opening(self, doc: Document, lesson_data: dict) -> None:
        """Add On Screen Opening visual directions."""
        on_screen = lesson_data.get("on_screen_opening")
        if not on_screen:
            return

        _add_styled_heading(doc, "On Screen (Opening)", level=3)

        if isinstance(on_screen, dict):
            for field_name in ["layout", "static_elements", "interactive_elements", "animation"]:
                val = on_screen.get(field_name)
                if val is not None:
                    if isinstance(val, list):
                        _add_label_value(doc, field_name.replace("_", " ").title(), ", ".join(str(v) for v in val))
                    else:
                        _add_label_value(doc, field_name.replace("_", " ").title(), _safe_str(val))

    def _add_narrated_explanation(self, doc: Document, lesson_data: dict) -> None:
        """Add Narrated Explanation section with concept subsections."""
        explanations = lesson_data.get("narrated_explanation")
        if not explanations:
            return

        _add_styled_heading(doc, "Narrated Explanation", level=2)

        if isinstance(explanations, list):
            for i, concept in enumerate(explanations, start=1):
                if not isinstance(concept, dict):
                    continue
                concept_name = _safe_str(concept.get("concept_name"), f"Concept {i}")
                bloom_level = _safe_str(concept.get("bloom_level"), "")

                heading_text = concept_name
                if bloom_level:
                    heading_text += f" [{bloom_level}]"
                _add_styled_heading(doc, heading_text, level=3)

                teacher_text = _safe_str(concept.get("teacher_explains"))
                if teacher_text:
                    _add_narration_line(doc, "Teacher explains", teacher_text)

                on_screen = concept.get("on_screen")
                if on_screen and isinstance(on_screen, dict):
                    _add_label_value(doc, "On Screen", str(on_screen))

                transition = _safe_str(concept.get("transition"))
                if transition:
                    para = doc.add_paragraph()
                    run = para.add_run(f"Transition: {transition}")
                    run.italic = True
                    run.font.size = Pt(10)

    def _add_interactive_activity(self, doc: Document, lesson_data: dict) -> None:
        """Add Interactive Activity section."""
        activity = lesson_data.get("interactive_activity")
        if not activity:
            return

        _add_styled_heading(doc, "Interactive Activity", level=2)

        if isinstance(activity, dict):
            act_type = _safe_str(activity.get("type"), "Not specified")
            bloom = _safe_str(activity.get("bloom_level"), "")
            instructions = _safe_str(activity.get("instructions"), "")

            _add_label_value(doc, "Type", act_type)
            if bloom:
                _add_label_value(doc, "Bloom's Level", bloom)
            if instructions:
                _add_label_value(doc, "Instructions", instructions)

            on_screen = activity.get("on_screen")
            if on_screen:
                _add_label_value(doc, "On Screen", str(on_screen))

            # Feedback hints
            for hint_key, hint_label in [
                ("feedback_hint_1", "Hint 1 (Gentle nudge)"),
                ("feedback_hint_2", "Hint 2 (More explicit)"),
                ("feedback_reveal", "Hint 3 (Reveal with reasoning)"),
            ]:
                hint = _safe_str(activity.get(hint_key))
                if hint:
                    _add_label_value(doc, hint_label, hint)

    def _add_doubts_discussion(self, doc: Document, lesson_data: dict) -> None:
        """Add Doubts & Discussion section."""
        doubts = lesson_data.get("doubts_discussion")
        if not doubts:
            return

        _add_styled_heading(doc, "Doubts & Discussion", level=2)

        if isinstance(doubts, list):
            for i, item in enumerate(doubts, start=1):
                if not isinstance(item, dict):
                    continue
                question = _safe_str(item.get("question"), "")
                bloom = _safe_str(item.get("bloom_level"), "")
                answer = _safe_str(item.get("answer"), "")
                clarification = _safe_str(item.get("teacher_clarification"), "")

                header = f"Q{i}"
                if bloom:
                    header += f" [{bloom}]"
                para = doc.add_paragraph()
                run = para.add_run(header + ": ")
                run.bold = True
                run.font.size = Pt(10)
                para.add_run(question).font.size = Pt(10)

                if answer:
                    _add_label_value(doc, "  Answer", answer)
                if clarification:
                    _add_label_value(doc, "  Teacher clarification", clarification)

    def _add_quick_quiz(self, doc: Document, lesson_data: dict) -> None:
        """Add Quick Quiz section with all questions."""
        quiz = lesson_data.get("quick_quiz")
        if not quiz:
            return

        _add_styled_heading(doc, f"Quick Quiz ({len(quiz)} Questions)", level=2)

        if isinstance(quiz, list):
            for q in quiz:
                if not isinstance(q, dict):
                    continue
                q_num = q.get("question_number", "?")
                q_type = _safe_str(q.get("type"), "")
                bloom = _safe_str(q.get("bloom_level"), "")
                prompt_text = _safe_str(q.get("prompt"), "")

                # Question header
                header_text = f"Question {q_num}"
                if q_type:
                    header_text += f" ({q_type})"
                if bloom:
                    header_text += f" [{bloom}]"

                para = doc.add_paragraph()
                run = para.add_run(header_text)
                run.bold = True
                run.font.size = Pt(11)

                if prompt_text:
                    doc.add_paragraph(prompt_text).style.font.size = Pt(10)

                # Options
                options = q.get("options")
                if options and isinstance(options, list):
                    for opt in options:
                        opt_para = doc.add_paragraph(style="List Bullet")
                        opt_para.add_run(_safe_str(opt)).font.size = Pt(10)

                # Answer
                answer = _safe_str(q.get("answer"))
                if answer:
                    _add_label_value(doc, "Answer", answer)

                # Feedback
                fc = _safe_str(q.get("feedback_correct"))
                fi = _safe_str(q.get("feedback_incorrect"))
                if fc:
                    _add_label_value(doc, "Feedback (Correct)", fc)
                if fi:
                    _add_label_value(doc, "Feedback (Incorrect)", fi)

                doc.add_paragraph("")  # spacing

    def _add_conclusion(self, doc: Document, lesson_data: dict) -> None:
        """Add Conclusion & Reflection section."""
        conclusion = lesson_data.get("conclusion")
        if not conclusion:
            return

        _add_styled_heading(doc, "Conclusion & Reflection", level=2)

        if isinstance(conclusion, dict):
            recap = _safe_str(conclusion.get("recap"))
            real_life = _safe_str(conclusion.get("real_life_connection"))
            reflection = _safe_str(conclusion.get("reflection_prompt"))

            if recap:
                _add_label_value(doc, "Recap", recap)
            if real_life:
                _add_label_value(doc, "Real-Life Connection", real_life)
            if reflection:
                _add_label_value(doc, "Reflection Prompt", reflection)
        elif isinstance(conclusion, str):
            doc.add_paragraph(conclusion)

    def _add_validation_appendix(self, doc: Document, report: dict) -> None:
        """Add validation report as an appendix."""
        doc.add_page_break()
        _add_styled_heading(doc, "APPENDIX: Validation Report", level=2)

        passed = report.get("passed", False)
        overall = report.get("overall_score", 0)

        status_text = "ALL CHECKS PASSED" if passed else "VALIDATION FAILED"
        para = doc.add_paragraph()
        run = para.add_run(f"Overall Status: {status_text}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) if passed else RGBColor(0xCC, 0x00, 0x00)

        _add_label_value(doc, "Overall Score", f"{overall}")

        checks = report.get("checks", [])
        if checks:
            _add_styled_heading(doc, "Individual Checks", level=3)
            for check in checks:
                name = check.get("name", "unknown")
                status = check.get("status", "unknown")
                icon = "PASS" if status == "passed" else ("WARN" if status == "warning" else "FAIL")
                _add_label_value(doc, f"  [{icon}] {name}", check.get("message", status))

        warnings = report.get("warnings", [])
        if warnings:
            _add_styled_heading(doc, "Warnings", level=3)
            for w in warnings:
                doc.add_paragraph(f"- {w.get('type', '')}: {w.get('message', '')}")

        errors = report.get("errors", [])
        if errors:
            _add_styled_heading(doc, "Errors", level=3)
            for e in errors:
                doc.add_paragraph(f"- {e.get('type', '')}: {e.get('message', '')}")
